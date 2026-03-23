from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .config_loader import config, FONT_NAME

def apply_font(run, font_config):
    run.font.name = FONT_NAME
    run.font.size = Pt(font_config['size_pt'])
    run.font.bold = font_config['bold']
    run.font.italic = font_config['italic']
    rgb = font_config['color_rgb']
    run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])

def setup_footer(doc):
    if not config['footer']['enabled']:
        return
    
    section = doc.sections[0]
    section.footer.is_linked_to_previous = False
    
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    
    cfg = config['fonts']['footer']
    
    run = p.add_run(config['footer']['left_symbol'])
    apply_font(run, cfg)
    
    run_page = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run_page._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run_page._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_page._r.append(fldChar2)
    
    apply_font(run_page, cfg)
    
    run = p.add_run(config['footer']['right_symbol'])
    apply_font(run, cfg)

def setup_styles(doc):
    h1 = doc.styles['Heading 1'].font
    h1.name = FONT_NAME
    h1.size = Pt(config['fonts']['conversation']['size_pt'])
    h1.bold = config['fonts']['conversation']['bold']
    h1.italic = config['fonts']['conversation']['italic']
    h1.color.rgb = RGBColor(*config['fonts']['conversation']['color_rgb'])
    
    h2 = doc.styles['Heading 2'].font
    h2.name = FONT_NAME
    h2.size = Pt(config['fonts']['chapter']['size_pt'])
    h2.bold = config['fonts']['chapter']['bold']
    h2.italic = config['fonts']['chapter']['italic']
    h2.color.rgb = RGBColor(*config['fonts']['chapter']['color_rgb'])

def add_table_of_contents(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run('Содержание')
    run.font.size = Pt(config['fonts']['text']['size_pt'])
    run.font.name = FONT_NAME
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    
    doc.add_page_break()