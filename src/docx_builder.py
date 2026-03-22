import re
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config_loader import config, FONT_NAME
from .docx_helpers import apply_font

def process_footnotes_in_text(element, text_config):
    """Обрабатывает элемент и возвращает список фрагментов с форматированием"""
    fragments = []
    
    def get_text_content(node):
        """Получает текстовое содержимое узла (рекурсивно)"""
        if isinstance(node, str):
            return node
        return node.get_text()
    
    def process_node(node, default_format='normal'):
        if isinstance(node, str):
            if node:
                fragments.append((node, default_format, None))
            return
        
        # Определяем формат для текущего узла
        current_format = default_format
        
        # Обработка <br> как переноса строки
        if node.name == 'br':
            fragments.append(('\n', current_format, None))
            return
        
        # Обработка <i> (курсив)
        if node.name == 'i':
            current_format = 'italic'
        
        # Обработка <a> со сноской
        if node.name == 'a' and node.get('href', '').startswith('#note'):
            note_text = node.get_text(strip=True)
            match = re.search(r'(\d+)', note_text)
            if match:
                note_number = match.group(1)
                fragments.append((note_number, 'superscript', None))
            return
        
        # Обработка <sup>
        if node.name == 'sup':
            for child in node.children:
                if child.name == 'a' and child.get('href', '').startswith('#note'):
                    note_text = child.get_text(strip=True)
                    match = re.search(r'(\d+)', note_text)
                    if match:
                        note_number = match.group(1)
                        fragments.append((note_number, 'superscript', None))
                else:
                    text = child if isinstance(child, str) else child.get_text()
                    if text:
                        fragments.append((text, 'superscript', None))
            return
        
        # Обработка <span> с цитатами
        if node.name == 'span':
            classes = node.get('class', [])
            if 'quote' in classes or 'synodal' in classes:
                current_format = 'italic'
            if 'church' in classes:
                current_format = 'italic'
        
        # Обработка <b> (жирный)
        if node.name == 'b':
            current_format = 'bold'
        
        # Рекурсивно обрабатываем детей
        children = list(node.children)
        for i, child in enumerate(children):
            process_node(child, current_format)
            
            # Проверяем, нужно ли добавить пробел между элементами
            if i < len(children) - 1:
                current_text = get_text_content(child).strip()
                next_text = get_text_content(children[i + 1]).strip()
                
                if current_text and next_text:
                    last_char = current_text[-1]
                    first_char = next_text[0]
                    
                    should_add_space = (
                        last_char.isalnum() and 
                        first_char.isalnum() and
                        last_char not in '.,!?;:)]»' and
                        first_char not in '.,!?;:([«'
                    )
                    
                    if should_add_space:
                        fragments.append((' ', current_format, None))
    
    process_node(element)
    return fragments

def add_heading_with_footnotes(doc, element, heading_level, font_config):
    if heading_level == 1:
        heading = doc.add_heading(level=1)
    else:
        heading = doc.add_heading(level=2)
    
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fragments = process_footnotes_in_text(element, font_config)
    
    for text, fmt, note_num in fragments:
        if not text:
            continue
        
        # Проверяем, есть ли в тексте переносы строк
        if '\n' in text:
            parts = text.split('\n')
            for i, part in enumerate(parts):
                if part:
                    run = heading.add_run(part)
                    apply_font(run, font_config)
                    if fmt == 'bold':
                        run.bold = True
                    elif fmt == 'italic':
                        run.italic = True
                    elif fmt == 'superscript':
                        run.font.superscript = True
                # После каждой части, кроме последней, добавляем перенос строки
                if i < len(parts) - 1:
                    heading.add_run().add_break()
        else:
            run = heading.add_run(text)
            apply_font(run, font_config)
            if fmt == 'bold':
                run.bold = True
            elif fmt == 'italic':
                run.italic = True
            elif fmt == 'superscript':
                run.font.superscript = True

def add_formatted_paragraph(doc, p_element, text_config):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(text_config.get('first_line_indent_cm', 0.76))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    fragments = process_footnotes_in_text(p_element, text_config)

    # Просто добавляем фрагменты как есть, без дополнительных пробелов
    for text, fmt, note_num in fragments:
        if not text:
            continue
        
        run = paragraph.add_run(text)
        apply_font(run, text_config)
        
        if fmt == 'bold':
            run.bold = True
        elif fmt == 'italic':
            run.italic = True
        elif fmt == 'superscript':
            run.font.superscript = True

def add_notes_section(doc, notes):
    if not notes:
        return
    
    p = doc.add_paragraph()
    run = p.add_run('★ ★ ★')
    run.font.size = Pt(config['fonts']['text']['size_pt'])
    run.font.name = FONT_NAME
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    h = doc.add_heading('Примечания', 2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        apply_font(run, config['fonts']['chapter'])
    
    for note in notes:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        
        run = p.add_run(f"{note['number']}")
        run.font.superscript = True
        apply_font(run, config['fonts']['text'])
        
        run = p.add_run(" ")
        apply_font(run, config['fonts']['text'])
        
        for fragment in note['fragments']:
            run = p.add_run(fragment['text'])
            apply_font(run, config['fonts']['text'])
            if fragment.get('italic'):
                run.italic = True
            if fragment.get('bold'):
                run.bold = True

def create_document(doc):
    """Создает и настраивает новый документ (заголовки, оглавление)"""
    from .docx_helpers import setup_footer, setup_styles, add_table_of_contents, apply_font
    
    setup_footer(doc)
    setup_styles(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_font(p.add_run(config['headers']['main_title']), config['fonts']['main_title'])
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_font(p.add_run(config['headers']['subtitle']), config['fonts']['subtitle'])
    
    doc.add_paragraph()
    add_table_of_contents(doc)