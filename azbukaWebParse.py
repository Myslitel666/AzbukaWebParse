import os

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src import (
    config,  # используется
    apply_font,  # используется
    add_heading_with_footnotes,  # используется
    add_formatted_paragraph,  # используется
    add_notes_section,  # используется
    get_conversations,  # используется
    fetch_all,  # используется
    create_document  # используется для настройки документа
)

# ========== ОСНОВНОЙ КОД ==========
def main():
    file_name = config['output_file']
    
    while True:
        if not os.path.exists(file_name):
            break
        try:
            os.remove(file_name)
            print(f'Старый файл "{file_name}" удалён')
            break
        except PermissionError:
            print(f'Файл "{file_name}" открыт!')
            print('Пожалуйста, закройте файл в Word и нажмите Enter...')
            input()
    
    print('\nСоздаю документ...')
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Cm(config['margins']['top_cm'])
    section.bottom_margin = Cm(config['margins']['bottom_cm'])
    section.left_margin = Cm(config['margins']['left_cm'])
    section.right_margin = Cm(config['margins']['right_cm'])
    
    # Настройка документа (футер, стили, заголовки, оглавление)
    create_document(doc)
    
    conversations = get_conversations()
    print(f"Найдено: {len(conversations)}")
    
    print("\nЗагрузка...")
    results = fetch_all(conversations, config['fonts']['text'])
    
    total = 0
    total_notes = 0
    
    for conv, chapters, is_fallback, notes in results:
        if conv['title'] and not conv.get('is_single_page', False):
            h = doc.add_heading(conv['title'], 1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                apply_font(run, config['fonts']['conversation'])
        
        if is_fallback:
            for ch in chapters:
                for p in ch['paragraphs']:
                    add_formatted_paragraph(doc, p, config['fonts']['text'])
                total += 1
        else:
            for ch in chapters:
                # Определяем уровень заголовка (2 или 3)
                level = ch.get('level', 2)
                if ch['element']:
                    add_heading_with_footnotes(doc, ch['element'], level, config['fonts']['chapter'])
                
                for p in ch['paragraphs']:
                    add_formatted_paragraph(doc, p, config['fonts']['text'])
                
                total += 1
        
        if notes:
            add_notes_section(doc, notes)
            total_notes += len(notes)
    
    doc.save(file_name)
    
    print(f"\nГотово: {file_name}")
    print(f"Глав: {total}")
    print(f"Примечаний: {total_notes}")
    
    try:
        os.startfile(file_name)
    except:
        pass

if __name__ == "__main__":
    main()